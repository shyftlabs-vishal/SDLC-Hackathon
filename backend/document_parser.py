"""Extract plain text from uploaded requirement documents."""

from __future__ import annotations

import io
from pathlib import Path

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_EXTRACTED_CHARS = 50_000
ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


class DocumentParseError(Exception):
    pass


def _validate_extension(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        allowed = ", ".join(sorted(ALLOWED_EXTENSIONS))
        raise DocumentParseError(
            f"Unsupported file type '{ext or '(none)'}'. Allowed: {allowed}"
        )
    return ext


def _validate_size(data: bytes) -> None:
    if not data:
        raise DocumentParseError("File is empty.")
    if len(data) > MAX_FILE_SIZE_BYTES:
        mb = len(data) / (1024 * 1024)
        raise DocumentParseError(f"File too large ({mb:.1f} MB). Maximum is 10 MB.")


def _extract_plain(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParseError("Could not decode text file.")


def _extract_pdf(data: bytes) -> str:
    if not data.startswith(b"%PDF"):
        raise DocumentParseError("File does not appear to be a valid PDF.")
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentParseError("PDF support is not installed on the server.") from exc

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text and page_text.strip():
            parts.append(page_text.strip())
    return "\n\n".join(parts)


def _extract_docx(data: bytes) -> str:
    if not data.startswith(b"PK"):
        raise DocumentParseError("File does not appear to be a valid DOCX document.")
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentParseError("DOCX support is not installed on the server.") from exc

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n\n".join(parts)


def extract_text_from_bytes(
    data: bytes,
    filename: str,
    *,
    content_type: str | None = None,
) -> dict[str, object]:
    del content_type  # extension + magic-byte checks are the source of truth
    ext = _validate_extension(filename)
    _validate_size(data)

    if ext in (".txt", ".md"):
        text = _extract_plain(data)
    elif ext == ".pdf":
        text = _extract_pdf(data)
    elif ext == ".docx":
        text = _extract_docx(data)
    else:
        raise DocumentParseError("Unsupported file type.")

    cleaned = text.strip()
    if not cleaned:
        raise DocumentParseError(
            "No readable text found in this file. "
            "Scanned or image-only PDFs are not supported."
        )

    truncated = len(cleaned) > MAX_EXTRACTED_CHARS
    if truncated:
        cleaned = cleaned[:MAX_EXTRACTED_CHARS]

    return {
        "text": cleaned,
        "filename": Path(filename).name,
        "char_count": len(cleaned),
        "truncated": truncated,
    }
